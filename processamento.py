import os
import uuid
import base64
import requests
import cohere
import fitz  # Importação correta do PyMuPDF
from docx import Document

from config import GOOGLE_API_KEY, COHERE_API_KEY

UPLOAD_FOLDER = 'uploads'
PROCESSED_FOLDER = 'processed'

def extrair_texto_com_api_key(image_path, api_key):
    """
    Realiza OCR na imagem utilizando a API do Google Cloud Vision.
    """
    with open(image_path, 'rb') as image_file:
        encoded_image = base64.b64encode(image_file.read()).decode()

    payload = {
        "requests": [
            {
                "image": {"content": encoded_image},
                "features": [{"type": "DOCUMENT_TEXT_DETECTION"}]
            }
        ]
    }

    url = f"https://vision.googleapis.com/v1/images:annotate?key={api_key}"
    response = requests.post(url, json=payload)

    if response.status_code == 200:
        data = response.json()
        try:
            return data['responses'][0]['fullTextAnnotation']['text']
        except KeyError:
            return "[Nenhum texto detectado]"
    else:
        raise Exception(f"Erro {response.status_code} na API Vision: {response.text}")

def extrair_texto_digital(file_path):
# Extração do texto em arquivos digitalizados
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".txt":
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    elif ext == ".docx":
        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs]
        return "\n".join(paragraphs)
    elif ext == ".pdf":
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text() + "\n"
        return text
    else:
        return ""

def split_text(text, max_chars=3000):
# Divide o texto em partes
    chunks = []
    while len(text) > max_chars:
        break_index = text.rfind('\n', 0, max_chars)
        if break_index == -1:
            break_index = max_chars
        chunks.append(text[:break_index])
        text = text[break_index:].strip()
    if text:
        chunks.append(text)
    return chunks

def corrigir_texto_com_cohere(texto, prompt):
# Envia o texto para a API da Cohere
    co = cohere.Client(COHERE_API_KEY)
    max_chars = 3000  # Limite para cada bloco 
    chunks = split_text(texto, max_chars)
    corrected_chunks = []
    
    for chunk in chunks:
        full_prompt = f"{prompt}\n\n{chunk}"
        response = co.generate(
            model='command-r-08-2024',  # Modificar, consultar o site do Cohere
            prompt=full_prompt,
            max_tokens=1000,
            temperature=0.3
        )
        corrected_chunks.append(response.generations[0].text.strip())
    
    return "\n".join(corrected_chunks)

def converter_pdf_para_imagens(pdf_path, output_folder):
# Converte o pdf para imagem
    doc = fitz.open(pdf_path)
    imagens = []
    
    for pagina in doc:
        pix = pagina.get_pixmap(dpi=200)
        nome_imagem = f"{uuid.uuid4()}_pagina_{pagina.number}.png"
        caminho_imagem = os.path.join(output_folder, nome_imagem)
        pix.save(caminho_imagem)
        imagens.append(caminho_imagem)
    
    return imagens

def processar_arquivo(file_path, prompt, mode):
# Processa o arquivo OCR ou puramente textual
    ext = os.path.splitext(file_path)[1].lower()
    texto_total = ""

    if mode == "ocr":
        if ext == ".pdf":
            # Converte PDF em imagens e aplica OCR a cada página
            imagens = converter_pdf_para_imagens(file_path, UPLOAD_FOLDER)
            for img in imagens:
                texto_total += extrair_texto_com_api_key(img, GOOGLE_API_KEY) + "\n"
        else:
            texto_total = extrair_texto_com_api_key(file_path, GOOGLE_API_KEY)
    elif mode == "texto":
        texto_total = extrair_texto_digital(file_path)
    else:
        raise ValueError("Modo inválido. Utilize 'ocr' ou 'texto'.")

    # Corrige todo o texto
    texto_corrigido = corrigir_texto_com_cohere(texto_total, prompt)
    
    base_nome = os.path.splitext(os.path.basename(file_path))[0]
    output_name = f"{base_nome}_corrigido.docx"
    output_path = os.path.join(PROCESSED_FOLDER, output_name)
    
    document = Document()
    document.add_paragraph(texto_corrigido)
    document.save(output_path)
    
    return output_name
